# #document_handler.py

# import os
# import docx
# import PyPDF2
# import streamlit as st

# def extract_text_from_document(file_path: str) -> str:
#     """
#     Extracts text from PDF and DOCX files.
#     Supports accessibility processing for transcription.
#     """
#     if not os.path.exists(file_path):
#         raise FileNotFoundError(f"File not found: {file_path}")

#     text = ""

#     # Handle PDF
#     if file_path.lower().endswith(".pdf"):
#         with open(file_path, "rb") as pdf_file:
#             reader = PyPDF2.PdfReader(pdf_file)
#             for page in reader.pages:
#                 text += page.extract_text() + "\n"

#     # Handle DOCX
#     elif file_path.lower().endswith(".docx"):
#         doc = docx.Document(file_path)
#         for para in doc.paragraphs:
#             text += para.text + "\n"

#     # Handle TXT
#     elif file_path.lower().endswith(".txt"):
#         with open(file_path, "r", encoding="utf-8") as f:
#             text = f.read()

#     else:
#         raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")

#     return text.strip()


#-------------------------------------------------------------------------------------------------------------------------

# document_handler.py

import os
import io
import docx
import PyPDF2
import streamlit as st

def extract_text_from_document(file_obj) -> str:
    """
    Extract text from PDF, DOCX, or TXT.
    Supports:
        - file path (str)
        - bytes or BytesIO (in-memory file)
    Returns: extracted text
    """
    # Determine if we have path or bytes
    if isinstance(file_obj, str):
        if not os.path.exists(file_obj):
            raise FileNotFoundError(f"File not found: {file_obj}")
        ext = os.path.splitext(file_obj)[1].lower()
        # For PDFs, DOCX, TXT we can use path directly
        file_bytes = None
        file_path = file_obj
    elif isinstance(file_obj, (bytes, bytearray, io.BytesIO)):
        # In-memory file
        if isinstance(file_obj, io.BytesIO):
            file_obj.seek(0)
            file_bytes = file_obj.read()
        else:
            file_bytes = file_obj
        # Need to detect extension if possible
        file_path = None
        ext = None  # will detect manually later
    else:
        raise ValueError("Unsupported input type. Provide file path or bytes/BytesIO.")

    text = ""

    # -------------------------
    # PDF
    # -------------------------
    if (file_path and file_path.lower().endswith(".pdf")) or (file_bytes and file_bytes[:4] == b"%PDF"):
        try:
            pdf_stream = open(file_path, "rb") if file_path else io.BytesIO(file_bytes)
            reader = PyPDF2.PdfReader(pdf_stream)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if file_path:
                pdf_stream.close()
        except Exception as e:
            raise RuntimeError(f"Failed to parse PDF: {e}")

    # -------------------------
    # DOCX
    # -------------------------
    elif (file_path and file_path.lower().endswith(".docx")) or (file_bytes and file_bytes[:2] == b"PK"):
        try:
            if file_path:
                doc = docx.Document(file_path)
            else:
                doc_stream = io.BytesIO(file_bytes)
                doc = docx.Document(doc_stream)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX: {e}")

    # -------------------------
    # TXT
    # -------------------------
    elif (file_path and file_path.lower().endswith(".txt")) or (file_bytes and b"\n" in file_bytes):
        try:
            if file_path:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            else:
                text = file_bytes.decode("utf-8", errors="ignore")
        except Exception as e:
            raise RuntimeError(f"Failed to parse TXT: {e}")

    else:
        raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")

    return text.strip()



def highlight_text_live(text: str, speed: float = 0.5):
    """
    Reads text line by line and highlights it while displaying.
    Useful for accessibility (screen reader simulation).
    """
    lines = text.split("\n")
    placeholder = st.empty()

    for i, line in enumerate(lines):
        highlighted = ""
        for j, l in enumerate(lines):
            if j == i:
                highlighted += f"<span style='background-color: yellow; color: black;'>{l}</span><br>"
            else:
                highlighted += f"{l}<br>"
        placeholder.markdown(highlighted, unsafe_allow_html=True)
        st.sleep(speed)  # wait before highlighting next line

