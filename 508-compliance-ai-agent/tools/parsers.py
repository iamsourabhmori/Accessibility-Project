# # tools/parsers.py

# from PyPDF2 import PdfReader
# from docx import Document
# from bs4 import BeautifulSoup

# def is_supported_file(filename: str) -> bool:
#     return any(filename.lower().endswith(ext) for ext in ['.pdf', '.docx', '.html', '.htm'])

# def parse_document(filepath: str):
#     if filepath.endswith('.pdf'):
#         return parse_pdf(filepath)
#     elif filepath.endswith('.docx'):
#         return parse_docx(filepath)
#     elif filepath.endswith('.html') or filepath.endswith('.htm'):
#         return parse_html_file(filepath)
#     else:
#         return None

# def parse_pdf(filepath: str):
#     text = ""
#     with open(filepath, "rb") as f:
#         reader = PdfReader(f)
#         for page in reader.pages:
#             text += page.extract_text()
#     return text

# def parse_docx(filepath: str):
#     doc = Document(filepath)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     return "\n".join(full_text)

# def parse_html_file(filepath: str):
#     with open(filepath, "r", encoding="utf-8") as f:
#         soup = BeautifulSoup(f, "html.parser")
#     return soup

#----------------------------------------------------------------------------------------------------------------------------

# tools/parsers.py

from PyPDF2 import PdfReader
from docx import Document
from bs4 import BeautifulSoup

def is_supported_file(filename: str) -> bool:
    return any(filename.lower().endswith(ext) for ext in ['.pdf', '.docx', '.html', '.htm'])

def parse_document(filepath: str):
    if filepath.lower().endswith('.pdf'):
        return parse_pdf(filepath)
    elif filepath.lower().endswith('.docx'):
        return parse_docx(filepath)
    elif filepath.lower().endswith(('.html', '.htm')):
        return parse_html_file(filepath)
    else:
        return None

def parse_pdf(filepath: str) -> str:
    """Extract text from PDF."""
    text = ""
    with open(filepath, "rb") as f:
        reader = PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()

def parse_docx(filepath: str) -> str:
    """Extract text from DOCX."""
    doc = Document(filepath)
    full_text = [para.text for para in doc.paragraphs if para.text.strip()]
    return "\n".join(full_text)

def parse_html_file(filepath: str) -> str:
    """Return HTML string from an HTML file."""
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        soup = BeautifulSoup(f, "html.parser")
    return str(soup)  # Now returns HTML as string
